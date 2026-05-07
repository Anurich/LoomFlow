"""DOCX loader → markdown.

Uses ``python-docx`` (lazy import). Walks the document body and
emits markdown:

* ``Heading 1`` / ``Heading 2`` / ... → ``#``, ``##``, ...
* Lists → ``-`` items
* Plain paragraphs → text
* Tables → markdown tables
"""

from __future__ import annotations

from pathlib import Path

from .base import Document


def _heading_level(style_name: str) -> int | None:
    """Return the heading level (1-6) for a Word style name, or None
    if it isn't a heading."""
    if not style_name:
        return None
    name = style_name.lower()
    if name == "title":
        return 1
    if name.startswith("heading "):
        try:
            level = int(name[len("heading ") :])
        except ValueError:
            return None
        return max(1, min(level, 6))
    return None


def _is_list_paragraph(paragraph: object) -> bool:
    """Best-effort detect of list paragraphs. Word's list styles
    aren't always reliable; this catches the common cases."""
    style = getattr(paragraph, "style", None)
    style_name = (
        getattr(style, "name", "") if style is not None else ""
    ).lower()
    return "list" in style_name or "bullet" in style_name


def _table_to_markdown(table: object) -> str:
    """Convert a ``docx.table.Table`` to a markdown table."""
    rows = []
    for row in getattr(table, "rows", []):
        cells = [
            (getattr(cell, "text", "") or "").replace("\n", " ").strip()
            for cell in getattr(row, "cells", [])
        ]
        cells = [c.replace("|", "\\|") for c in cells]
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    out = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in rows[1:]:
        padded = row + [""] * (len(header) - len(row))
        out.append("| " + " | ".join(padded) + " |")
    return "\n".join(out)


def load_docx(path: str | Path) -> Document:
    """Load a ``.docx`` file → markdown.

    Requires ``python-docx``:
    ``pip install 'jeevesagent[loader-docx]'``.
    """
    try:
        import docx  # type: ignore[import-not-found, import-untyped]
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "python-docx is not installed. "
            "Install with: pip install 'jeevesagent[loader-docx]'."
        ) from exc

    p = Path(path)
    document = docx.Document(str(p))

    # Walk the document in order, interleaving paragraphs and tables.
    # python-docx stores them in document.element.body.iter() — we
    # use the simpler high-level API and process paragraphs +
    # tables in document order using their _element index.
    paragraphs = list(document.paragraphs)
    tables = list(document.tables)

    # Build a map from xml element to (kind, payload) so we can walk
    # the body in order.
    body = document.element.body
    parts: list[str] = []
    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            # Find the matching paragraph
            for para in paragraphs:
                if para._element is child:  # noqa: SLF001
                    text = (para.text or "").strip()
                    if not text:
                        parts.append("")
                        continue
                    level = _heading_level(
                        para.style.name if para.style else ""
                    )
                    if level is not None:
                        parts.append(f"{'#' * level} {text}")
                    elif _is_list_paragraph(para):
                        parts.append(f"- {text}")
                    else:
                        parts.append(text)
                    break
        elif tag == "tbl":
            for table in tables:
                if table._element is child:  # noqa: SLF001
                    md = _table_to_markdown(table)
                    if md:
                        parts.append("")
                        parts.append(md)
                        parts.append("")
                    break

    # Collapse runs of blank lines.
    out: list[str] = []
    prev_blank = False
    for line in parts:
        is_blank = not line.strip()
        if is_blank and prev_blank:
            continue
        out.append(line)
        prev_blank = is_blank
    content = "\n".join(out).strip() + "\n"

    # Document properties
    core_props = document.core_properties
    title = (core_props.title or "").strip() if core_props else ""
    if not title and p.stem:
        title = p.stem

    if not content.startswith("# "):
        content = f"# {title}\n\n{content}"

    return Document(
        content=content,
        metadata={
            "source": str(p),
            "format": "docx",
            "title": title,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        },
    )
