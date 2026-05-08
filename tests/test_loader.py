"""Tests for document loaders.

Each loader is tested with a small fixture file created in tmp_path
on the fly. PDF / DOCX / Excel / HTML tests use real format files
(generated via their producer libs, which we already require for
the `loader` extras in pyproject).
"""

from __future__ import annotations

from pathlib import Path

import pytest

from jeevesagent.loader import (
    Document,
    load,
    load_csv,
    load_docx,
    load_excel,
    load_html,
    load_markdown,
    load_pdf,
    load_text,
    load_tsv,
)

# ---------------------------------------------------------------------------
# text + markdown
# ---------------------------------------------------------------------------


def test_load_markdown_returns_content_unchanged(tmp_path: Path) -> None:
    f = tmp_path / "doc.md"
    f.write_text("# Title\n\nSome **markdown** content.\n")
    doc = load_markdown(f)
    assert isinstance(doc, Document)
    assert "# Title" in doc.content
    assert "**markdown**" in doc.content
    assert doc.metadata["format"] == "md"
    assert doc.metadata["source"] == str(f)


def test_load_text_wraps_in_markdown_heading(tmp_path: Path) -> None:
    f = tmp_path / "notes.txt"
    f.write_text("Line one.\nLine two.\n")
    doc = load_text(f)
    assert doc.content.startswith("# notes.txt")
    assert "Line one." in doc.content
    assert doc.metadata["format"] == "txt"


# ---------------------------------------------------------------------------
# CSV / TSV
# ---------------------------------------------------------------------------


def test_load_csv_produces_markdown_table(tmp_path: Path) -> None:
    f = tmp_path / "data.csv"
    f.write_text(
        "name,age,city\n"
        "Alice,30,Tokyo\n"
        "Bob,25,Paris\n"
    )
    doc = load_csv(f)
    assert "| name | age | city |" in doc.content
    assert "| --- | --- | --- |" in doc.content
    assert "| Alice | 30 | Tokyo |" in doc.content
    assert doc.metadata["format"] == "csv"
    assert doc.metadata["row_count"] == 2  # excluding header
    assert doc.metadata["column_count"] == 3


def test_load_csv_escapes_pipe_characters(tmp_path: Path) -> None:
    f = tmp_path / "weird.csv"
    f.write_text("col\nfoo|bar\n")
    doc = load_csv(f)
    # The pipe should be escaped so the markdown table stays valid.
    assert "foo\\|bar" in doc.content


def test_load_tsv_uses_tabs(tmp_path: Path) -> None:
    f = tmp_path / "data.tsv"
    f.write_text("a\tb\tc\n1\t2\t3\n")
    doc = load_tsv(f)
    assert "| a | b | c |" in doc.content
    assert "| 1 | 2 | 3 |" in doc.content
    assert doc.metadata["format"] == "tsv"


def test_load_csv_pads_short_rows(tmp_path: Path) -> None:
    f = tmp_path / "short.csv"
    f.write_text("a,b,c\n1,2\n")  # second row missing one column
    doc = load_csv(f)
    # Padded with empty cell so the table stays valid
    assert "| 1 | 2 |  |" in doc.content


def test_load_csv_empty_file(tmp_path: Path) -> None:
    f = tmp_path / "empty.csv"
    f.write_text("")
    doc = load_csv(f)
    assert "(empty)" in doc.content


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _make_pdf(path: Path, pages: list[str], *, title: str = "Test PDF Title") -> None:
    """Create a small text-bearing PDF for testing.

    Uses ``reportlab`` so the PDF actually carries an extractable
    text layer — important now that the loader is backed by
    ``unstructured`` and we want to assert on real content, not
    just structural metadata.
    """
    pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=LETTER)
    c.setTitle(title)
    width, height = LETTER
    for body in pages:
        text = c.beginText(72, height - 72)
        text.setFont("Helvetica", 11)
        for line in body.split("\n"):
            text.textLine(line)
        c.drawText(text)
        c.showPage()
    c.save()


def test_load_pdf_extracts_per_page_content(tmp_path: Path) -> None:
    """The default ``unstructured`` backend should produce a
    markdown ``Document`` with a ``# title``, per-page ``## Page N``
    sections, and the actual extracted text from EVERY page —
    including the last. The pypdf-backed loader silently dropped
    pages on extraction errors, which was the source of the
    lower-half-of-PDF retrieval miss."""
    f = tmp_path / "small.pdf"
    _make_pdf(
        f,
        pages=[
            "Founding facts.\nAcme was founded in 2008 in Berlin.",
            "Operating facts.\nAcme's CEO is Mira Castellanos.",
        ],
        title="Acme Handbook",
    )
    doc = load_pdf(f)

    assert isinstance(doc, Document)
    assert doc.metadata["format"] == "pdf"
    assert doc.metadata["page_count"] == 2
    assert doc.metadata["backend"] == "unstructured"
    assert doc.metadata["strategy"] == "fast"
    assert "## Page 1" in doc.content
    assert "## Page 2" in doc.content
    assert "Acme was founded in 2008" in doc.content
    assert "Mira Castellanos" in doc.content
    assert "(no extractable text)" not in doc.content


def test_load_pdf_rejects_unknown_backend(tmp_path: Path) -> None:
    f = tmp_path / "x.pdf"
    _make_pdf(f, pages=["hi"], title="x")
    with pytest.raises(ValueError, match="unknown backend"):
        load_pdf(f, backend="bogus")


def test_load_pdf_raises_for_missing_file(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        load_pdf(tmp_path / "does_not_exist.pdf")


def test_load_pdf_docling_backend_extracts_content(tmp_path: Path) -> None:
    """The docling backend produces the same ``Document`` shape as
    the unstructured one — same ``page_count``, same ``title``,
    same per-page sections — but uses Docling's ML-based
    extraction underneath. Skipped when docling isn't installed,
    since it brings heavy ML deps."""
    pytest.importorskip("docling")
    f = tmp_path / "small.pdf"
    _make_pdf(
        f,
        pages=[
            "Founding facts.\nAcme was founded in 2008 in Berlin.",
            "Operating facts.\nAcme's CEO is Mira Castellanos.",
        ],
        title="Acme Handbook",
    )
    doc = load_pdf(f, backend="docling")

    assert isinstance(doc, Document)
    assert doc.metadata["format"] == "pdf"
    assert doc.metadata["backend"] == "docling"
    assert doc.metadata["page_count"] >= 1
    # Same content contract: real text from every page surfaces.
    assert "Acme was founded in 2008" in doc.content
    assert "Mira Castellanos" in doc.content


# ---------------------------------------------------------------------------
# Regression: "lower-half pages have no answers" symptom
# ---------------------------------------------------------------------------
#
# The historical pypdf-backed loader silently swallowed per-page
# extraction errors with ``except Exception: text = ""``, producing
# the symptom where questions about content near the END of a PDF
# went unanswered while content from the start surfaced cleanly.
# These tests build an 8-page PDF with deterministic, distinctive
# content per page and assert every page's marker survives the
# round-trip — proving both backends fix the regression.


_LOWER_HALF_PAGES: list[tuple[str, str]] = [
    ("PAGE_ONE_MARKER",   "On page one we mention CONST_ALPHA explicitly."),
    ("PAGE_TWO_MARKER",   "On page two we mention CONST_BRAVO explicitly."),
    ("PAGE_THREE_MARKER", "On page three we mention CONST_CHARLIE explicitly."),
    ("PAGE_FOUR_MARKER",  "On page four we mention CONST_DELTA explicitly."),
    ("PAGE_FIVE_MARKER",  "On page five we mention CONST_ECHO explicitly."),
    ("PAGE_SIX_MARKER",   "On page six we mention CONST_FOXTROT explicitly."),
    ("PAGE_SEVEN_MARKER", "On page seven we mention CONST_GOLF explicitly."),
    ("PAGE_EIGHT_MARKER", "On page eight we mention CONST_HOTEL explicitly."),
]


def _make_eight_page_pdf(path: Path) -> None:
    """Build the 8-page deterministic-content PDF used by the
    lower-half regression tests."""
    bodies = [
        f"{marker}\n\n{sentence}\n\n"
        f"This is filler so the page has body text.\n"
        f"Search phrase: find_me_on_{marker.lower()}."
        for marker, sentence in _LOWER_HALF_PAGES
    ]
    _make_pdf(path, bodies, title="Eight Page Reference")


def test_load_pdf_unstructured_recovers_every_page(tmp_path: Path) -> None:
    """Regression test for the "lower-half pages drop content"
    symptom. Every page's unique marker AND its unique constant
    name MUST appear in the loaded markdown. If pypdf-style silent
    drops creep back in, this fails with a clear "PAGE_X_MARKER
    missing" assertion."""
    f = tmp_path / "eight.pdf"
    _make_eight_page_pdf(f)
    doc = load_pdf(f, backend="unstructured")

    assert doc.metadata["page_count"] == 8
    # Every page-section header is present in arrival order.
    for n in range(1, 9):
        assert f"## Page {n}" in doc.content, f"missing section header for page {n}"
    # Every per-page marker AND its unique constant survive.
    for marker, sentence in _LOWER_HALF_PAGES:
        const = sentence.split("CONST_")[1].split(" ")[0]
        assert marker in doc.content, f"{marker} missing from extracted content"
        assert f"CONST_{const}" in doc.content, (
            f"CONST_{const} missing — page content was dropped"
        )
    # No silent-empty placeholders.
    assert "(no extractable text)" not in doc.content


def test_load_pdf_docling_recovers_every_page(tmp_path: Path) -> None:
    """Same regression contract under the docling backend. Skipped
    when docling isn't installed, since the deps are heavy."""
    pytest.importorskip("docling")
    f = tmp_path / "eight.pdf"
    _make_eight_page_pdf(f)
    doc = load_pdf(f, backend="docling")

    # Docling's page_count comes from the ``DoclingDocument``
    # object; both 8-and-positive accepted (some docling versions
    # expose num_pages differently, but the content assertions
    # below are the contract that matters).
    assert doc.metadata["page_count"] >= 1
    for marker, sentence in _LOWER_HALF_PAGES:
        const = sentence.split("CONST_")[1].split(" ")[0]
        assert marker in doc.content, (
            f"{marker} missing from docling-extracted content"
        )
        assert f"CONST_{const}" in doc.content, (
            f"CONST_{const} missing under docling — page content dropped"
        )


def test_load_pdf_chunks_cover_all_pages(tmp_path: Path) -> None:
    """End-to-end RAG-shaped check: load → chunk → search. The
    user-reported symptom was that questions about page 8 content
    came back blank. After the loader fix, a recursive chunker at
    600/80 (the user's reported config) emits chunks covering
    every page, including the last."""
    from jeevesagent.loader.chunking import RecursiveChunker

    f = tmp_path / "eight.pdf"
    _make_eight_page_pdf(f)
    doc = load_pdf(f, backend="unstructured")

    chunker = RecursiveChunker(chunk_size=600, chunk_overlap=80)
    chunks = chunker.split(doc.content, source=str(f))

    # At least one chunk must contain each page's unique constant —
    # otherwise the retriever would never see that page's content.
    joined = " || ".join(c.content for c in chunks)
    for _marker, sentence in _LOWER_HALF_PAGES:
        const = sentence.split("CONST_")[1].split(" ")[0]
        assert f"CONST_{const}" in joined, (
            f"chunker dropped page content (CONST_{const} not in any chunk)"
        )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def _make_docx(
    path: Path, title: str, sections: list[tuple[int, str]]
) -> None:
    """Create a small .docx with a title + headings + paragraphs."""
    pytest.importorskip("docx")
    import docx

    document = docx.Document()
    document.core_properties.title = title
    for level, text in sections:
        if level == 0:
            document.add_paragraph(text)
        else:
            document.add_heading(text, level=level)
    document.save(str(path))


def test_load_docx_preserves_headings_and_paragraphs(
    tmp_path: Path,
) -> None:
    f = tmp_path / "doc.docx"
    _make_docx(
        f,
        title="My Doc",
        sections=[
            (1, "Introduction"),
            (0, "This is the intro paragraph."),
            (2, "Details"),
            (0, "Some details here."),
        ],
    )
    doc = load_docx(f)
    assert doc.metadata["format"] == "docx"
    assert "# Introduction" in doc.content
    assert "## Details" in doc.content
    assert "intro paragraph" in doc.content
    assert "details here" in doc.content


def test_load_docx_handles_tables(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    import docx

    f = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_heading("Report", level=1)
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "92"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "87"
    document.save(str(f))

    doc = load_docx(f)
    assert "| Name | Score |" in doc.content
    assert "| Alice | 92 |" in doc.content


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------


def test_load_excel_each_sheet_becomes_section(tmp_path: Path) -> None:
    pytest.importorskip("openpyxl")
    import openpyxl

    f = tmp_path / "data.xlsx"
    wb = openpyxl.Workbook()
    sheet1 = wb.active
    sheet1.title = "Sales"
    sheet1.append(["Q1", "Q2", "Q3"])
    sheet1.append([100, 200, 300])
    sheet2 = wb.create_sheet("Costs")
    sheet2.append(["Item", "Cost"])
    sheet2.append(["Servers", 5000])
    wb.save(str(f))

    doc = load_excel(f)
    assert doc.metadata["format"] == "xlsx"
    assert doc.metadata["sheet_count"] == 2
    assert doc.metadata["sheet_names"] == ["Sales", "Costs"]
    assert "## Sales" in doc.content
    assert "## Costs" in doc.content
    assert "| Q1 | Q2 | Q3 |" in doc.content
    assert "| 100 | 200 | 300 |" in doc.content
    assert "| Servers | 5000 |" in doc.content


def test_load_excel_trims_empty_trailing_rows(tmp_path: Path) -> None:
    pytest.importorskip("openpyxl")
    import openpyxl

    f = tmp_path / "padded.xlsx"
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "Data"
    sheet.append(["a", "b"])
    sheet.append([1, 2])
    # Excel sometimes leaves blank trailing rows; loader should trim.
    sheet.append([None, None])
    sheet.append([None, None])
    wb.save(str(f))

    doc = load_excel(f)
    # Should NOT have empty rows in the markdown table
    lines = [line for line in doc.content.splitlines() if "| " in line]
    # Header + separator + 1 data row = 3 lines (no empty rows after)
    data_rows = [
        line for line in lines if not line.startswith("| --- ")
    ]
    assert len([line for line in data_rows if "1" in line]) == 1


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def test_load_html_extracts_headings_and_paragraphs(
    tmp_path: Path,
) -> None:
    pytest.importorskip("bs4")
    f = tmp_path / "page.html"
    f.write_text(
        "<html><head><title>My Page</title></head>"
        "<body>"
        "<h1>Big Title</h1>"
        "<p>First paragraph.</p>"
        "<h2>Subsection</h2>"
        "<p>Second paragraph.</p>"
        "<ul><li>Item one</li><li>Item two</li></ul>"
        "<script>console.log('ignore me');</script>"
        "</body></html>"
    )
    doc = load_html(f)
    assert doc.metadata["format"] == "html"
    assert doc.metadata["title"] == "My Page"
    assert "# My Page" in doc.content
    assert "# Big Title" in doc.content
    assert "## Subsection" in doc.content
    assert "First paragraph." in doc.content
    assert "- Item one" in doc.content
    # Script content is dropped
    assert "console.log" not in doc.content


def test_load_html_extracts_tables(tmp_path: Path) -> None:
    pytest.importorskip("bs4")
    f = tmp_path / "table.html"
    f.write_text(
        "<table>"
        "<tr><th>Name</th><th>Age</th></tr>"
        "<tr><td>Alice</td><td>30</td></tr>"
        "</table>"
    )
    doc = load_html(f)
    assert "| Name | Age |" in doc.content
    assert "| Alice | 30 |" in doc.content


# ---------------------------------------------------------------------------
# dispatch.load — auto-detect by extension
# ---------------------------------------------------------------------------


def test_load_dispatches_md(tmp_path: Path) -> None:
    f = tmp_path / "doc.md"
    f.write_text("# Hello")
    doc = load(f)
    assert doc.metadata["format"] == "md"


def test_load_dispatches_csv(tmp_path: Path) -> None:
    f = tmp_path / "data.csv"
    f.write_text("a,b\n1,2\n")
    doc = load(f)
    assert doc.metadata["format"] == "csv"


def test_load_dispatches_xlsx(tmp_path: Path) -> None:
    pytest.importorskip("openpyxl")
    import openpyxl

    f = tmp_path / "wb.xlsx"
    wb = openpyxl.Workbook()
    wb.active["A1"] = "hi"
    wb.save(str(f))
    doc = load(f)
    assert doc.metadata["format"] == "xlsx"


def test_load_rejects_unknown_extension(tmp_path: Path) -> None:
    f = tmp_path / "weird.xyz"
    f.write_text("nope")
    with pytest.raises(ValueError, match="unsupported extension"):
        load(f)


def test_load_dispatches_html(tmp_path: Path) -> None:
    pytest.importorskip("bs4")
    f = tmp_path / "page.html"
    f.write_text("<html><body><p>hi</p></body></html>")
    doc = load(f)
    assert doc.metadata["format"] == "html"
